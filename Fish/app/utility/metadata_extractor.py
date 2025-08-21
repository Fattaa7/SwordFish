import os
import hashlib
import mimetypes
from typing import Dict, Any, Optional
import PyPDF2
from docx import Document as DocxDocument
import chardet
import json
from pathlib import Path
import io


class MetadataExtractor:
    """Extract metadata from various file types"""
    
    @staticmethod
    def extract_metadata(file_path: str, file_content: Optional[bytes] = None) -> Dict[str, Any]:
        """
        Extract metadata from a file
        
        Args:
            file_path: Path to the file
            file_content: Optional file content as bytes (for uploaded files)
            
        Returns:
            Dictionary containing extracted metadata
        """
        try:
            # Get file info
            file_info = MetadataExtractor._get_file_info(file_path, file_content)
            
            # Extract content-based metadata
            content_meta = MetadataExtractor._extract_content_metadata(file_path, file_content)
            
            # Combine all metadata
            metadata = {
                "title": content_meta.get("title", file_info["filename"]),
                "language": content_meta.get("language", "unknown"),
                "meta": {
                    "author": content_meta.get("author"),
                    "pages": content_meta.get("pages"),
                    "tags": content_meta.get("tags", []),
                    "file_size": file_info["size"],
                    "mime_type": file_info["mime_type"],
                    "checksum": file_info["checksum"]
                }
            }
            
            return metadata
            
        except Exception as e:
            # Return basic metadata if extraction fails
            return {
                "title": os.path.basename(file_path),
                "language": "unknown",
                "meta": {
                    "author": None,
                    "pages": None,
                    "tags": [],
                    "error": str(e)
                }
            }
    
    @staticmethod
    def _get_file_info(file_path: str, file_content: Optional[bytes] = None) -> Dict[str, Any]:
        """Get basic file information"""
        if file_content:
            size = len(file_content)
            checksum = hashlib.sha256(file_content).hexdigest()
        else:
            size = os.path.getsize(file_path)
            with open(file_path, 'rb') as f:
                checksum = hashlib.sha256(f.read()).hexdigest()
        
        mime_type, _ = mimetypes.guess_type(file_path)
        if not mime_type:
            mime_type = "application/octet-stream"
        
        return {
            "filename": os.path.basename(file_path),
            "size": size,
            "mime_type": mime_type,
            "checksum": checksum
        }
    
    @staticmethod
    def _extract_content_metadata(file_path: str, file_content: Optional[bytes] = None) -> Dict[str, Any]:
        """Extract metadata from file content based on file type"""
        mime_type, _ = mimetypes.guess_type(file_path)
        
        if mime_type == "application/pdf":
            return MetadataExtractor._extract_pdf_metadata(file_path, file_content)
        elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"]:
            return MetadataExtractor._extract_docx_metadata(file_path, file_content)
        elif mime_type in ["text/plain", "text/html", "text/csv"]:
            return MetadataExtractor._extract_text_metadata(file_path, file_content)
        else:
            return {}
    
    @staticmethod
    def _extract_pdf_metadata(file_path: str, file_content: Optional[bytes] = None) -> Dict[str, Any]:
        """Extract metadata from PDF files"""
        try:
            print(f"Extracting PDF metadata from {file_path}")
            print(f"File content provided: {file_content is not None}")
            
            pdf_reader = PyPDF2.PdfReader(file_path)
            info = pdf_reader.metadata
            num_pages = len(pdf_reader.pages)

            print(f"PDF Metadata: {info}")
            print(f"Number of pages: {num_pages}")
            text = ""
            title = None
            if num_pages > 0:
                first_page = pdf_reader.pages[0]
                text = first_page.extract_text() or ""
                if text:
                    lines = [line.strip() for line in text.split("\n") if line.strip()]
                    if lines:
                        title = lines[0][:100]

            return {
                "title": info.title or title or os.path.basename(file_path),
                "author": info.author,
                "pages": num_pages,
                "tags": [],
                "language": MetadataExtractor._detect_language_from_text(text) if text else "unknown"
            }
        except Exception as e:
            return {
                "title": os.path.basename(file_path),
                "author": None,
                "pages": None,
                "tags": [],
                "language": "unknown",
                "error": str(e)
            }
    
    @staticmethod
    def _extract_docx_metadata(file_path: str, file_content: Optional[bytes] = None) -> Dict[str, Any]:
        """Extract metadata from DOCX files"""
        try:
            if file_content:
                doc = DocxDocument(io.BytesIO(file_content))
            else:
                doc = DocxDocument(file_path)
            
            # Extract core properties
            core_props = doc.core_properties
            title = core_props.title or os.path.basename(file_path)
            author = core_props.author
            
            # Count paragraphs as a rough page estimate
            paragraph_count = len(doc.paragraphs)
            estimated_pages = max(1, paragraph_count // 20)  # Rough estimate
            
            # Extract text for language detection
            text = "\n".join([p.text for p in doc.paragraphs if p.text])
            
            return {
                "title": title,
                "author": author,
                "pages": estimated_pages,
                "tags": [],
                "language": MetadataExtractor._detect_language_from_text(text) if text else "unknown"
            }
        except Exception as e:
            return {
                "title": os.path.basename(file_path),
                "author": None,
                "pages": None,
                "tags": [],
                "language": "unknown",
                "error": str(e)
            }
    
    @staticmethod
    def _extract_text_metadata(file_path: str, file_content: Optional[bytes] = None) -> Dict[str, Any]:
        """Extract metadata from text files"""
        try:
            if file_content:
                text = file_content.decode('utf-8', errors='ignore')
            else:
                # Try to detect encoding
                with open(file_path, 'rb') as f:
                    raw_data = f.read()
                    detected = chardet.detect(raw_data)
                    encoding = detected['encoding'] or 'utf-8'
                    text = raw_data.decode(encoding, errors='ignore')
            
            # Simple title extraction: first non-empty line
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            title = lines[0][:100] if lines else os.path.basename(file_path)
            
            # Language detection
            language = MetadataExtractor._detect_language_from_text(text)
            
            return {
                "title": title,
                "author": None,
                "pages": max(1, len(lines) // 50),  # Rough page estimate
                "tags": [],
                "language": language
            }
        except Exception as e:
            return {
                "title": os.path.basename(file_path),
                "author": None,
                "pages": None,
                "tags": [],
                "language": "unknown",
                "error": str(e)
            }
    
    @staticmethod
    def _detect_language_from_text(text: str) -> str:
        """Simple language detection based on common words"""
        if not text:
            return "unknown"
        
        # Simple heuristics for common languages
        text_lower = text.lower()
        
        # English
        english_words = ['the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by']
        if any(word in text_lower for word in english_words):
            return "en"
        
        # French
        french_words = ['le', 'la', 'les', 'et', 'ou', 'mais', 'dans', 'sur', 'avec', 'par']
        if any(word in text_lower for word in french_words):
            return "fr"
        
        # Spanish
        spanish_words = ['el', 'la', 'los', 'las', 'y', 'o', 'pero', 'en', 'con', 'por']
        if any(word in text_lower for word in spanish_words):
            return "es"
        
        # German
        german_words = ['der', 'die', 'das', 'und', 'oder', 'aber', 'in', 'auf', 'mit', 'von']
        if any(word in text_lower for word in german_words):
            return "de"
        
        return "unknown"
